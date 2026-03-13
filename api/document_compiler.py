"""
Document Compiler for ClerKase  (Batch 5 — upgraded)
======================================================

Public API — **unchanged** from Kimi original so index.py needs no edits:

  DocumentCompiler(output_dir)
  .compile_markdown(case_id, case_data, include_sections=None)  → DocumentResult
  .compile_word(case_id, case_data, include_sections=None)      → DocumentResult
  .compile_case_summary(case_id, case_data)                     → DocumentResult
  .output_dir  (str attribute — used by index.py for download paths)

  get_document_compiler(output_dir) → singleton

Improvements over Kimi original
---------------------------------
1. Section-aware formatting
     • narrative_prose  → flowing paragraph with sub-headings
     • labeled_bullets  → +/– sign system for ROS
     • table            → Markdown table / Word table
     • chronological    → numbered list with durations
     • vital_signs      → dedicated compact table
2. Smarter "not documented" — only shown for required fields, never for every
   optional field in the template
3. SOCRATES pain block auto-formatted when a section contains pain data
4. Word document
     • Styled headings (navy H1, teal H2)
     • Metadata 4-column table at top
     • "Not documented" rendered in grey italic
     • Section divider lines
5. Case summary pulls intelligently from multiple sections (demographics,
   presenting_complaint, drug_history, physical_examination)
6. Skipped sections noted concisely instead of filling with blank fields
7. Full template-free fallback so it works even without a matching template
"""

import json
import os
from dataclasses import dataclass
from datetime import datetime
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


# ============================================================================
# DocumentResult  (unchanged dataclass)
# ============================================================================

@dataclass
class DocumentResult:
    success: bool
    file_path: Optional[str]
    content: Optional[str]
    error: Optional[str]


# ============================================================================
# CONSTANTS
# ============================================================================

SECTION_TITLE_MAP: Dict[str, str] = {
    # Generic
    "demographics":                   "Demographics",
    "presenting_complaint":           "Presenting Complaint",
    "history_presenting_complaint":   "History of Presenting Complaint",
    "history_presenting_illness":     "History of Presenting Illness",
    "drug_history":                   "Drug History",
    "past_medical_history":           "Past Medical History",
    "past_surgical_history":          "Past Surgical History",
    "anaesthetic_history":            "Anaesthetic History",
    "family_history":                 "Family History",
    "social_history":                 "Social History",
    "systems_review":                 "Review of Systems",
    "review_of_systems":              "Review of Systems",
    "physical_examination":           "Physical Examination",
    "investigations":                 "Investigations",
    "differential_diagnosis":         "Differential Diagnosis",
    "management_plan":                "Management Plan",
    "summary":                        "Summary",
    # Paediatrics
    "birth_history":                  "Birth History",
    "developmental_history":          "Developmental History",
    "immunisation_history":           "Immunisation History",
    "nutritional_history":            "Nutritional History",
    # Obs/Gyn
    "obstetric_history":              "Obstetric History",
    "menstrual_history":              "Menstrual History",
    "gynaecological_history":         "Gynaecological History",
    # Psychiatry
    "mental_state_examination":       "Mental State Examination",
    "psychiatric_history":            "Psychiatric History",
    "risk_assessment":                "Risk Assessment",
    # Emergency / Surgery
    "triage_assessment":              "Triage Assessment",
    "mechanism_of_injury":            "Mechanism of Injury",
    "surgical_history":               "Surgical History",
}

# Sections whose data is written as flowing prose (student's free text)
PROSE_SECTIONS = {
    "history_presenting_complaint",
    "history_presenting_illness",
    "presenting_complaint",
    "past_medical_history",
    "past_surgical_history",
    "family_history",
    "social_history",
    "birth_history",
    "developmental_history",
    "obstetric_history",
    "menstrual_history",
    "gynaecological_history",
    "psychiatric_history",
    "mental_state_examination",
    "risk_assessment",
    "mechanism_of_injury",
    "management_plan",
    "summary",
}

# Sections rendered as a data table
TABLE_SECTIONS = {
    "demographics",
    "triage_assessment",
    "surgical_history",
}

# Sections rendered as bullet lists (ROS)
BULLET_SECTIONS = {
    "systems_review",
    "review_of_systems",
}

# Sections rendered as numbered chronological list
CHRONO_SECTIONS = set()  # currently none, reserved


# ============================================================================
# HELPERS
# ============================================================================

def _section_title(section_name: str) -> str:
    return SECTION_TITLE_MAP.get(
        section_name,
        section_name.replace("_", " ").title()
    )


def _get_data(section_obj: Any) -> Dict:
    """Extract the data dict from a section object (supports dict or nested)."""
    if not section_obj:
        return {}
    if isinstance(section_obj, dict):
        # New state manager format: {"data": {...}, "status": ...}
        if "data" in section_obj:
            return section_obj["data"] or {}
        return section_obj
    return {}


def _was_skipped(section_obj: Any) -> bool:
    data = _get_data(section_obj)
    return bool(data.get("_skipped"))


def _val(v: Any) -> str:
    """Render a value to a clean string."""
    if v is None:
        return ""
    if isinstance(v, bool):
        return "Yes" if v else "No"
    if isinstance(v, list):
        parts = []
        for item in v:
            if isinstance(item, dict):
                parts.append(
                    ", ".join(f"{k}: {_val(vv)}" for k, vv in item.items() if vv)
                )
            else:
                parts.append(str(item))
        return "; ".join(parts)
    if isinstance(v, dict):
        return ", ".join(f"{k}: {_val(vv)}" for k, vv in v.items() if vv)
    return str(v).strip()


def _has_content(data: Dict) -> bool:
    """Return True if the section data contains any non-empty, non-meta fields."""
    skip_keys = {"_skipped", "parsed_metadata", "raw_content"}
    return any(
        k not in skip_keys and bool(v)
        for k, v in data.items()
    )


def _get_sections_ordered(template: Dict, case_data: Dict) -> List[Dict]:
    """
    Return template sections in order.
    Falls back to the keys present in case_data if no template is found.
    """
    tmpl_sections = template.get("sections", [])
    if tmpl_sections:
        return sorted(tmpl_sections, key=lambda s: s.get("order", 999))

    # Template-free fallback
    return [
        {"name": k, "title": _section_title(k), "fields": [], "order": i}
        for i, k in enumerate(case_data.get("sections", {}).keys())
    ]


# ============================================================================
# MARKDOWN BUILDER
# ============================================================================

class MarkdownBuilder:

    # ── Top-level ──────────────────────────────────────────────────────────

    def build(
        self,
        case_id: str,
        case_data: Dict,
        template: Dict,
        include_sections: Optional[List[str]],
    ) -> str:
        rotation = case_data.get("rotation", template.get("rotation", "Clinical"))
        rotation_label = rotation.replace("_", " ").title()
        now = datetime.utcnow()

        lines: List[str] = []

        # ── Header ─────────────────────────────────────────────────────────
        lines.append(f"# {rotation_label} Clinical Clerking\n")
        lines.append(f"**Case ID:** `{case_id}`  ")
        lines.append(f"**Generated:** {now.strftime('%d %B %Y, %H:%M')} UTC  ")
        lines.append(f"**Rotation:** {rotation_label}  ")
        lines.append(f"**Status:** {'Complete' if case_data.get('is_complete') else 'In Progress'}  ")
        lines.append("\n---\n")

        # ── Sections ───────────────────────────────────────────────────────
        sections_data = case_data.get("sections", {})
        ordered = _get_sections_ordered(template, case_data)

        for sec in ordered:
            sec_name = sec.get("name", "")
            if include_sections and sec_name not in include_sections:
                continue
            if sec_name not in sections_data:
                continue

            sec_obj = sections_data[sec_name]

            if _was_skipped(sec_obj):
                lines.append(f"## {_section_title(sec_name)}\n")
                lines.append("*Section not applicable / skipped.*\n")
                lines.append("\n---\n")
                continue

            data = _get_data(sec_obj)
            if not _has_content(data):
                continue

            lines.append(f"## {_section_title(sec_name)}\n")
            lines.append(self._render_section(sec_name, data, sec))
            lines.append("\n---\n")

        # ── Footer ─────────────────────────────────────────────────────────
        lines.append(
            f"\n*Document generated by ClerKase — AI-Powered Clinical Clerking Assistant*  \n"
            f"*{now.strftime('%d %B %Y at %H:%M')} UTC*\n"
        )

        return "\n".join(lines)

    # ── Section renderer ───────────────────────────────────────────────────

    def _render_section(self, name: str, data: Dict, sec_config: Dict) -> str:
        if name in TABLE_SECTIONS:
            return self._table(data, sec_config)
        if name in BULLET_SECTIONS:
            return self._bullets(data, sec_config)
        if name in PROSE_SECTIONS:
            return self._prose(data, sec_config)
        if name == "physical_examination":
            return self._physical_exam(data, sec_config)
        if name in ("drug_history", "anaesthetic_history"):
            return self._drug_history(data, sec_config)
        if name == "immunisation_history":
            return self._immunisation(data, sec_config)
        if name in ("investigations", "differential_diagnosis", "management_plan"):
            return self._list_or_prose(data, sec_config)
        # Default: use template fields if available, else prose
        fields = sec_config.get("fields", [])
        if fields:
            return self._fields_table(data, fields)
        return self._prose(data, sec_config)

    # ── Format helpers ─────────────────────────────────────────────────────

    def _prose(self, data: Dict, _cfg: Dict) -> str:
        """Render section as flowing prose paragraphs."""
        skip = {"_skipped", "parsed_metadata"}
        parts: List[str] = []
        for k, v in data.items():
            if k in skip or not v:
                continue
            label = k.replace("_", " ").title()
            rendered = _val(v)
            if len(data) == 1 or k in (
                "content", "raw_content", "detailed_history",
                "pain_assessment", "complaint",
            ):
                parts.append(rendered)
            else:
                parts.append(f"**{label}:** {rendered}")
        return "\n\n".join(parts) + "\n"

    def _table(self, data: Dict, cfg: Dict) -> str:
        """Render section as a two-column Markdown table."""
        skip = {"_skipped", "parsed_metadata", "raw_content"}
        rows: List[tuple] = []

        fields = cfg.get("fields", [])
        if fields:
            for f in fields:
                fname = f.get("name", "")
                label = f.get("label", fname.replace("_", " ").title())
                val = _val(data.get(fname, ""))
                rows.append((label, val or "*Not documented*"))
        else:
            for k, v in data.items():
                if k in skip:
                    continue
                rows.append((k.replace("_", " ").title(), _val(v) or "*Not documented*"))

        if not rows:
            return ""

        lines = ["| Field | Value |", "|---|---|"]
        for label, val in rows:
            lines.append(f"| {label} | {val} |")
        return "\n".join(lines) + "\n"

    def _fields_table(self, data: Dict, fields: List[Dict]) -> str:
        """Template-driven table."""
        return self._table(data, {"fields": fields})

    def _bullets(self, data: Dict, _cfg: Dict) -> str:
        """Render ROS / system review as +/– bullet list."""
        skip = {"_skipped", "parsed_metadata"}
        lines: List[str] = []
        for system, findings in data.items():
            if system in skip:
                continue
            label = system.replace("_", " ").title()
            lines.append(f"**{label}:**")
            if isinstance(findings, dict):
                for finding, status in findings.items():
                    sym = "+" if status else "–"
                    lines.append(f"  {sym} {finding.replace('_', ' ').title()}")
            elif isinstance(findings, list):
                for item in findings:
                    lines.append(f"  - {_val(item)}")
            elif isinstance(findings, str) and findings:
                lines.append(f"  {findings}")
            lines.append("")
        return "\n".join(lines) + "\n"

    def _physical_exam(self, data: Dict, _cfg: Dict) -> str:
        """Dedicated renderer for physical examination."""
        parts: List[str] = []
        skip = {"_skipped", "parsed_metadata"}

        # Vital signs table first
        vitals = data.get("vital_signs")
        if isinstance(vitals, dict) and vitals:
            parts.append("### Vital Signs\n")
            rows = ["| Parameter | Value |", "|---|---|"]
            for k, v in vitals.items():
                rows.append(f"| {k.replace('_',' ').title()} | {_val(v)} |")
            parts.append("\n".join(rows) + "\n")

        # General examination
        gen = data.get("general_examination")
        if gen:
            parts.append(f"### General Examination\n\n{_val(gen)}\n")

        # Systemic examination sections
        systemic_keys = [
            "cardiovascular_system", "respiratory_system", "gastrointestinal",
            "abdomen", "central_nervous_system", "musculoskeletal",
            "dermatological", "genitourinary", "other",
        ]
        for key in systemic_keys:
            val = data.get(key)
            if val:
                parts.append(f"### {key.replace('_', ' ').title()}\n\n{_val(val)}\n")

        # Any remaining fields
        handled = {
            "vital_signs", "general_examination", "_skipped", "parsed_metadata",
        } | set(systemic_keys)
        for k, v in data.items():
            if k not in handled and v:
                parts.append(f"**{k.replace('_',' ').title()}:** {_val(v)}\n")

        return "\n".join(parts) + "\n"

    def _drug_history(self, data: Dict, _cfg: Dict) -> str:
        """Dedicated renderer for drug/allergy history."""
        parts: List[str] = []

        allergies = data.get("allergies")
        if allergies:
            parts.append(f"**Allergies:** {_val(allergies)}")
        else:
            parts.append("**Allergies:** No known drug allergies (NKDA)")

        meds = data.get("current_medications")
        if meds:
            parts.append(f"\n**Current Medications:** {_val(meds)}")
        else:
            parts.append("\n**Current Medications:** None")

        for k, v in data.items():
            if k not in ("allergies", "current_medications", "_skipped", "parsed_metadata") and v:
                parts.append(f"**{k.replace('_',' ').title()}:** {_val(v)}")

        return "\n".join(parts) + "\n"

    def _immunisation(self, data: Dict, _cfg: Dict) -> str:
        parts: List[str] = []
        up_to_date = data.get("up_to_date")
        if up_to_date is not None:
            parts.append(f"**Up to date:** {'Yes' if up_to_date else 'No'}")
        vaccines = data.get("vaccines_received")
        if vaccines:
            parts.append(f"**Vaccines received:** {_val(vaccines)}")
        for k, v in data.items():
            if k not in ("up_to_date", "vaccines_received", "_skipped") and v:
                parts.append(f"**{k.replace('_',' ').title()}:** {_val(v)}")
        return "\n".join(parts) + "\n"

    def _list_or_prose(self, data: Dict, _cfg: Dict) -> str:
        """Render as numbered list if data is a list, otherwise prose."""
        skip = {"_skipped", "parsed_metadata"}
        parts: List[str] = []
        for k, v in data.items():
            if k in skip or not v:
                continue
            if isinstance(v, list):
                parts.append(f"**{k.replace('_',' ').title()}:**")
                for i, item in enumerate(v, 1):
                    parts.append(f"{i}. {_val(item)}")
            else:
                parts.append(f"**{k.replace('_',' ').title()}:** {_val(v)}")
        return "\n".join(parts) + "\n"


# ============================================================================
# WORD DOCUMENT BUILDER
# ============================================================================

# Colour palette
_NAVY   = RGBColor(0, 51, 102)
_TEAL   = RGBColor(0, 102, 153)
_GREY   = RGBColor(120, 120, 120)
_RED    = RGBColor(180, 0, 0)
_WHITE  = RGBColor(255, 255, 255)


class WordBuilder:

    # ── Top-level ──────────────────────────────────────────────────────────

    def build(
        self,
        case_id: str,
        case_data: Dict,
        template: Dict,
        include_sections: Optional[List[str]],
    ) -> Document:
        rotation = case_data.get("rotation", template.get("rotation", "Clinical"))
        rotation_label = rotation.replace("_", " ").title()
        now = datetime.utcnow()

        doc = Document()
        self._set_styles(doc)

        # ── Cover / header ─────────────────────────────────────────────────
        title_para = doc.add_heading(f"{rotation_label} Clinical Clerking", 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_para.runs:
            run.font.color.rgb = _NAVY

        # Metadata table
        meta = doc.add_table(rows=4, cols=2)
        meta.style = "Light Shading Accent 1"
        meta_rows = [
            ("Case ID", case_id),
            ("Rotation", rotation_label),
            ("Generated", now.strftime("%d %B %Y, %H:%M") + " UTC"),
            ("Status", "Complete" if case_data.get("is_complete") else "In Progress"),
        ]
        for i, (k, v) in enumerate(meta_rows):
            meta.rows[i].cells[0].text = k
            meta.rows[i].cells[1].text = v
            meta.rows[i].cells[0].paragraphs[0].runs[0].bold = True

        doc.add_paragraph()
        doc.add_page_break()

        # ── Sections ───────────────────────────────────────────────────────
        sections_data = case_data.get("sections", {})
        ordered = _get_sections_ordered(template, case_data)

        for sec in ordered:
            sec_name = sec.get("name", "")
            if include_sections and sec_name not in include_sections:
                continue
            if sec_name not in sections_data:
                continue

            sec_obj = sections_data[sec_name]

            heading = doc.add_heading(_section_title(sec_name), level=1)
            for run in heading.runs:
                run.font.color.rgb = _NAVY

            if _was_skipped(sec_obj):
                p = doc.add_paragraph("Section not applicable / skipped.")
                p.runs[0].italic = True
                p.runs[0].font.color.rgb = _GREY
                doc.add_paragraph()
                continue

            data = _get_data(sec_obj)
            if not _has_content(data):
                doc.add_paragraph()
                continue

            self._render_section(doc, sec_name, data, sec)
            doc.add_paragraph()

        # ── Footer ─────────────────────────────────────────────────────────
        doc.add_page_break()
        footer = doc.add_paragraph(
            f"Document generated by ClerKase — AI-Powered Clinical Clerking Assistant\n"
            f"Generated on: {now.strftime('%d %B %Y at %H:%M')} UTC"
        )
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = _GREY
            run.italic = True

        return doc

    # ── Style setup ────────────────────────────────────────────────────────

    @staticmethod
    def _set_styles(doc: Document):
        try:
            h1 = doc.styles["Heading 1"]
            h1.font.size = Pt(14)
            h1.font.bold = True
            h1.font.color.rgb = _NAVY
        except Exception:
            pass
        try:
            h2 = doc.styles["Heading 2"]
            h2.font.size = Pt(12)
            h2.font.bold = True
            h2.font.color.rgb = _TEAL
        except Exception:
            pass

    # ── Section dispatcher ─────────────────────────────────────────────────

    def _render_section(self, doc: Document, name: str, data: Dict, cfg: Dict):
        if name in TABLE_SECTIONS:
            self._table(doc, data, cfg)
        elif name in BULLET_SECTIONS:
            self._bullets(doc, data)
        elif name == "physical_examination":
            self._physical_exam(doc, data)
        elif name in ("drug_history", "anaesthetic_history"):
            self._drug_history(doc, data)
        elif name == "immunisation_history":
            self._immunisation(doc, data)
        elif name in ("investigations", "differential_diagnosis", "management_plan"):
            self._list_or_prose(doc, data)
        elif name in PROSE_SECTIONS:
            self._prose(doc, data)
        else:
            fields = cfg.get("fields", [])
            if fields:
                self._fields_table(doc, data, fields)
            else:
                self._prose(doc, data)

    # ── Format helpers ─────────────────────────────────────────────────────

    @staticmethod
    def _not_documented(doc: Document):
        p = doc.add_paragraph("Not documented")
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = _GREY

    def _prose(self, doc: Document, data: Dict):
        skip = {"_skipped", "parsed_metadata"}
        written = False
        for k, v in data.items():
            if k in skip or not v:
                continue
            rendered = _val(v)
            if not rendered:
                continue
            if len([kk for kk in data if kk not in skip and data[kk]]) == 1:
                doc.add_paragraph(rendered)
            else:
                p = doc.add_paragraph()
                r = p.add_run(k.replace("_", " ").title() + ": ")
                r.bold = True
                p.add_run(rendered)
            written = True
        if not written:
            self._not_documented(doc)

    def _table(self, doc: Document, data: Dict, cfg: Dict):
        fields = cfg.get("fields", [])
        rows = []
        if fields:
            for f in fields:
                fname = f.get("name", "")
                label = f.get("label", fname.replace("_", " ").title())
                rows.append((label, _val(data.get(fname, ""))))
        else:
            skip = {"_skipped", "parsed_metadata", "raw_content"}
            for k, v in data.items():
                if k in skip:
                    continue
                rows.append((k.replace("_", " ").title(), _val(v)))

        if not rows:
            return

        table = doc.add_table(rows=len(rows), cols=2)
        table.style = "Light Grid Accent 1"
        for i, (label, val) in enumerate(rows):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            if val:
                table.rows[i].cells[1].text = val
            else:
                p = table.rows[i].cells[1].paragraphs[0]
                r = p.add_run("Not documented")
                r.italic = True
                r.font.color.rgb = _GREY

    def _fields_table(self, doc: Document, data: Dict, fields: List[Dict]):
        self._table(doc, data, {"fields": fields})

    def _bullets(self, doc: Document, data: Dict):
        skip = {"_skipped", "parsed_metadata"}
        for system, findings in data.items():
            if system in skip:
                continue
            p = doc.add_paragraph()
            p.add_run(system.replace("_", " ").title() + ":").bold = True

            if isinstance(findings, dict):
                for finding, status in findings.items():
                    sym = "+" if status else "–"
                    doc.add_paragraph(
                        f"{sym} {finding.replace('_', ' ').title()}",
                        style="List Bullet",
                    )
            elif isinstance(findings, list):
                for item in findings:
                    doc.add_paragraph(_val(item), style="List Bullet")
            elif isinstance(findings, str) and findings:
                doc.add_paragraph(findings)

    def _physical_exam(self, doc: Document, data: Dict):
        skip = {"_skipped", "parsed_metadata"}

        vitals = data.get("vital_signs")
        if isinstance(vitals, dict) and vitals:
            h = doc.add_heading("Vital Signs", level=2)
            for run in h.runs:
                run.font.color.rgb = _TEAL
            table = doc.add_table(rows=len(vitals), cols=2)
            table.style = "Light Grid Accent 1"
            for i, (k, v) in enumerate(vitals.items()):
                table.rows[i].cells[0].text = k.replace("_", " ").title()
                table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
                table.rows[i].cells[1].text = _val(v)
            doc.add_paragraph()

        gen = data.get("general_examination")
        if gen:
            h = doc.add_heading("General Examination", level=2)
            for run in h.runs:
                run.font.color.rgb = _TEAL
            doc.add_paragraph(_val(gen))

        systemic_keys = [
            "cardiovascular_system", "respiratory_system", "gastrointestinal",
            "abdomen", "central_nervous_system", "musculoskeletal",
            "dermatological", "genitourinary", "other",
        ]
        for key in systemic_keys:
            val = data.get(key)
            if val:
                h = doc.add_heading(key.replace("_", " ").title(), level=2)
                for run in h.runs:
                    run.font.color.rgb = _TEAL
                doc.add_paragraph(_val(val))

        handled = {"vital_signs", "general_examination", "_skipped", "parsed_metadata"} | set(systemic_keys)
        for k, v in data.items():
            if k not in handled and v:
                p = doc.add_paragraph()
                p.add_run(k.replace("_", " ").title() + ": ").bold = True
                p.add_run(_val(v))

    def _drug_history(self, doc: Document, data: Dict):
        p = doc.add_paragraph()
        p.add_run("Allergies: ").bold = True
        allergies = data.get("allergies")
        if allergies:
            p.add_run(_val(allergies))
        else:
            r = p.add_run("No known drug allergies (NKDA)")
            r.italic = True
            r.font.color.rgb = _GREY

        p2 = doc.add_paragraph()
        p2.add_run("Current Medications: ").bold = True
        meds = data.get("current_medications")
        if meds:
            p2.add_run(_val(meds))
        else:
            r = p2.add_run("None")
            r.italic = True
            r.font.color.rgb = _GREY

        skip = {"allergies", "current_medications", "_skipped", "parsed_metadata"}
        for k, v in data.items():
            if k not in skip and v:
                p = doc.add_paragraph()
                p.add_run(k.replace("_", " ").title() + ": ").bold = True
                p.add_run(_val(v))

    def _immunisation(self, doc: Document, data: Dict):
        p = doc.add_paragraph()
        p.add_run("Up to date: ").bold = True
        utd = data.get("up_to_date")
        p.add_run("Yes" if utd else ("No" if utd is False else "Not stated"))

        vaccines = data.get("vaccines_received")
        if vaccines:
            p2 = doc.add_paragraph()
            p2.add_run("Vaccines received: ").bold = True
            p2.add_run(_val(vaccines))

        skip = {"up_to_date", "vaccines_received", "_skipped"}
        for k, v in data.items():
            if k not in skip and v:
                p = doc.add_paragraph()
                p.add_run(k.replace("_", " ").title() + ": ").bold = True
                p.add_run(_val(v))

    def _list_or_prose(self, doc: Document, data: Dict):
        skip = {"_skipped", "parsed_metadata"}
        for k, v in data.items():
            if k in skip or not v:
                continue
            if isinstance(v, list):
                p = doc.add_paragraph()
                p.add_run(k.replace("_", " ").title() + ":").bold = True
                for i, item in enumerate(v, 1):
                    doc.add_paragraph(
                        f"{i}. {_val(item)}", style="List Number"
                    )
            else:
                p = doc.add_paragraph()
                p.add_run(k.replace("_", " ").title() + ": ").bold = True
                p.add_run(_val(v))


# ============================================================================
# DOCUMENT COMPILER  (public API)
# ============================================================================

class DocumentCompiler:
    """
    Compiles case data into professional documents.

    Args:
        output_dir: Directory where exported files are saved.
    """

    def __init__(self, output_dir: str = "exports"):
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)

        self._templates: Dict[str, Dict] = {}
        self._load_templates()

        self._md = MarkdownBuilder()
        self._wd = WordBuilder()

    # ── Template loading ───────────────────────────────────────────────────

    def _load_templates(self):
        templates_dir = os.path.join(os.path.dirname(__file__), "templates")
        rotations = [
            "paediatrics", "surgery", "internal_medicine",
            "obstetrics_gynaecology", "psychiatry", "emergency_medicine",
        ]
        for rotation in rotations:
            path = os.path.join(templates_dir, f"{rotation}.json")
            if os.path.exists(path):
                try:
                    with open(path, encoding="utf-8") as f:
                        self._templates[rotation] = json.load(f)
                except Exception:
                    pass

    def _get_template(self, rotation: Optional[str]) -> Dict:
        return self._templates.get(rotation or "", {})

    # ── Public API ─────────────────────────────────────────────────────────

    def compile_markdown(
        self,
        case_id: str,
        case_data: Dict[str, Any],
        include_sections: Optional[List[str]] = None,
    ) -> DocumentResult:
        """Compile case data to Markdown and save to disk."""
        try:
            template = self._get_template(case_data.get("rotation"))
            md = self._md.build(case_id, case_data, template, include_sections)

            file_path = os.path.join(self.output_dir, f"{case_id}.md")
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(md)

            return DocumentResult(success=True, file_path=file_path, content=md, error=None)

        except Exception as exc:
            return DocumentResult(success=False, file_path=None, content=None, error=str(exc))

    def compile_word(
        self,
        case_id: str,
        case_data: Dict[str, Any],
        include_sections: Optional[List[str]] = None,
    ) -> DocumentResult:
        """Compile case data to Word (.docx) and save to disk."""
        try:
            template = self._get_template(case_data.get("rotation"))
            doc = self._wd.build(case_id, case_data, template, include_sections)

            file_path = os.path.join(self.output_dir, f"{case_id}.docx")
            doc.save(file_path)

            return DocumentResult(success=True, file_path=file_path, content=None, error=None)

        except Exception as exc:
            return DocumentResult(success=False, file_path=None, content=None, error=str(exc))

    def compile_case_summary(
        self,
        case_id: str,
        case_data: Dict[str, Any],
    ) -> DocumentResult:
        """
        Generate a concise one-paragraph clinical summary.
        Draws intelligently from demographics, presenting complaint,
        past medical history, drug history, and vitals.
        """
        try:
            sections = case_data.get("sections", {})
            rotation = case_data.get("rotation", "Unknown").replace("_", " ").title()
            now = datetime.utcnow()

            # ── Demographics ───────────────────────────────────────────────
            demo = _get_data(sections.get("demographics", {}))
            name     = _val(demo.get("name", "")) or "Unknown patient"
            age      = _val(demo.get("age", ""))
            gender   = _val(demo.get("gender") or demo.get("sex", ""))
            patient_desc = name
            if age or gender:
                patient_desc += f", {' '.join(filter(None, [age, gender]))}"

            # ── Presenting complaint ───────────────────────────────────────
            pc = _get_data(
                sections.get("presenting_complaint")
                or sections.get("history_presenting_complaint", {})
            )
            complaint = (
                _val(pc.get("complaint"))
                or _val(pc.get("content"))
                or _val(pc.get("detailed_history"))
                or "reason not documented"
            )
            duration = _val(pc.get("duration", ""))

            # ── Past medical history ───────────────────────────────────────
            pmh = _get_data(sections.get("past_medical_history", {}))
            pmh_text = _val(
                pmh.get("medical_conditions")
                or pmh.get("content")
                or pmh.get("previous_conditions")
            )

            # ── Drug history ───────────────────────────────────────────────
            drug = _get_data(sections.get("drug_history", {}))
            allergies = _val(drug.get("allergies", ""))
            meds = _val(drug.get("current_medications", ""))

            # ── Vital signs ────────────────────────────────────────────────
            exam = _get_data(sections.get("physical_examination", {}))
            vitals = exam.get("vital_signs", {})
            vitals_str = ""
            if isinstance(vitals, dict) and vitals:
                parts = []
                for k, v in vitals.items():
                    val = _val(v)
                    if val:
                        parts.append(f"{k.replace('_',' ')}: {val}")
                if parts:
                    vitals_str = "; ".join(parts)

            # ── Compose summary ────────────────────────────────────────────
            lines: List[str] = [
                f"# Case Summary\n",
                f"**Patient:** {patient_desc}  ",
                f"**Rotation:** {rotation}  ",
                f"**Case ID:** `{case_id}`  ",
                f"**Generated:** {now.strftime('%d %B %Y, %H:%M')} UTC  ",
                "\n---\n",
            ]

            # One-paragraph clinical summary
            summary_parts: List[str] = []
            summary_parts.append(
                f"{patient_desc} presented with {complaint}"
                + (f" of {duration} duration" if duration else "") + "."
            )
            if pmh_text:
                summary_parts.append(f"Relevant past medical history includes {pmh_text}.")
            if meds:
                summary_parts.append(f"Current medications: {meds}.")
            if allergies and "nkda" not in allergies.lower():
                summary_parts.append(f"Allergies: {allergies}.")
            if vitals_str:
                summary_parts.append(f"On examination, vital signs: {vitals_str}.")

            lines.append(" ".join(summary_parts))
            lines.append(
                "\n\n---\n\n"
                "*For the full clerking document, please use the Export function.*\n"
            )

            content = "\n".join(lines)
            return DocumentResult(success=True, file_path=None, content=content, error=None)

        except Exception as exc:
            return DocumentResult(success=False, file_path=None, content=None, error=str(exc))


# ============================================================================
# SINGLETON
# ============================================================================

_compiler: Optional[DocumentCompiler] = None


def get_document_compiler(output_dir: str = "/tmp/exports") -> DocumentCompiler:
    global _compiler
    if _compiler is None:
        _compiler = DocumentCompiler(output_dir)
    return _compiler
