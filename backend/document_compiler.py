"""
Document Compiler for ClerKase
Generates Markdown and Word (.docx) documents from case data
"""

import os
import json
from typing import Dict, List, Optional, Any
from datetime import datetime
from dataclasses import dataclass

# Document generation libraries
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


@dataclass
class DocumentResult:
    """Result from document compilation"""
    success: bool
    file_path: Optional[str]
    content: Optional[str]
    error: Optional[str]


class DocumentCompiler:
    """
    Compiles case data into professional documents
    """
    
    def __init__(self, output_dir: str = "exports"):
        self.output_dir = output_dir
        self._ensure_output_dir()
        self._templates = {}
        self._load_templates()
    
    def _ensure_output_dir(self):
        """Ensure output directory exists"""
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
    
    def _load_templates(self):
        """Load rotation templates"""
        templates_dir = os.path.join(os.path.dirname(__file__), "templates")
        
        rotations = [
            "paediatrics",
            "surgery",
            "internal_medicine",
            "obstetrics_gynaecology",
            "psychiatry",
            "emergency_medicine"
        ]
        
        for rotation in rotations:
            template_path = os.path.join(templates_dir, f"{rotation}.json")
            if os.path.exists(template_path):
                with open(template_path, 'r') as f:
                    self._templates[rotation] = json.load(f)
    
    def get_template(self, rotation: str) -> Optional[Dict]:
        """Get template for a rotation"""
        return self._templates.get(rotation)
    
    def compile_markdown(
        self,
        case_id: str,
        case_data: Dict[str, Any],
        include_sections: Optional[List[str]] = None
    ) -> DocumentResult:
        """
        Compile case data to Markdown
        
        Args:
            case_id: The case ID
            case_data: Complete case data
            include_sections: Optional list of sections to include
            
        Returns:
            DocumentResult with markdown content
        """
        try:
            template = self.get_template(case_data.get("rotation"))
            if not template:
                return DocumentResult(
                    success=False,
                    file_path=None,
                    content=None,
                    error=f"Template not found for rotation: {case_data.get('rotation')}"
                )
            
            # Build markdown content
            md_content = self._build_markdown(case_id, case_data, template, include_sections)
            
            # Save to file
            file_path = os.path.join(self.output_dir, f"{case_id}.md")
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(md_content)
            
            return DocumentResult(
                success=True,
                file_path=file_path,
                content=md_content,
                error=None
            )
            
        except Exception as e:
            return DocumentResult(
                success=False,
                file_path=None,
                content=None,
                error=str(e)
            )
    
    def compile_word(
        self,
        case_id: str,
        case_data: Dict[str, Any],
        include_sections: Optional[List[str]] = None
    ) -> DocumentResult:
        """
        Compile case data to Word document
        
        Args:
            case_id: The case ID
            case_data: Complete case data
            include_sections: Optional list of sections to include
            
        Returns:
            DocumentResult with file path
        """
        try:
            template = self.get_template(case_data.get("rotation"))
            if not template:
                return DocumentResult(
                    success=False,
                    file_path=None,
                    content=None,
                    error=f"Template not found for rotation: {case_data.get('rotation')}"
                )
            
            # Build Word document
            doc = self._build_word_document(case_id, case_data, template, include_sections)
            
            # Save to file
            file_path = os.path.join(self.output_dir, f"{case_id}.docx")
            doc.save(file_path)
            
            return DocumentResult(
                success=True,
                file_path=file_path,
                content=None,
                error=None
            )
            
        except Exception as e:
            return DocumentResult(
                success=False,
                file_path=None,
                content=None,
                error=str(e)
            )
    
    def _build_markdown(
        self,
        case_id: str,
        case_data: Dict[str, Any],
        template: Dict,
        include_sections: Optional[List[str]]
    ) -> str:
        """Build Markdown content"""
        
        rotation_name = template.get("rotation", "Unknown").replace("_", " ").title()
        
        md = f"""# Clinical Clerking: {rotation_name}

**Case ID:** {case_id}  
**Generated:** {datetime.utcnow().strftime('%Y-%m-%d %H:%M')} UTC  
**Rotation:** {rotation_name}  

---

"""
        
        # Get sections in order
        sections = sorted(template.get("sections", []), key=lambda s: s.get("order", 999))
        
        for section in sections:
            section_name = section.get("name")
            
            # Skip if not in include list
            if include_sections and section_name not in include_sections:
                continue
            
            section_title = section.get("title", section_name.replace("_", " ").title())
            section_data = case_data.get("sections", {}).get(section_name, {})
            
            md += f"## {section_title}\n\n"
            
            # Add section content
            for field in section.get("fields", []):
                field_name = field.get("name")
                field_label = field.get("label", field_name.replace("_", " ").title())
                field_value = section_data.get("data", {}).get(field_name, "")
                
                if field_value:
                    md += f"**{field_label}:** {field_value}\n\n"
                else:
                    md += f"**{field_label}:** *Not documented*\n\n"
            
            md += "\n---\n\n"
        
        # Add footer
        md += """---

*Document generated by ClerKase - AI-Powered Clinical Clerking Assistant*
"""
        
        return md
    
    def _build_word_document(
        self,
        case_id: str,
        case_data: Dict[str, Any],
        template: Dict,
        include_sections: Optional[List[str]]
    ) -> Document:
        """Build Word document"""
        
        doc = Document()
        
        rotation_name = template.get("rotation", "Unknown").replace("_", " ").title()
        
        # Title
        title = doc.add_heading(f'Clinical Clerking: {rotation_name}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_paragraph(f'Case ID: {case_id}')
        doc.add_paragraph(f'Generated: {datetime.utcnow().strftime("%Y-%m-%d %H:%M")} UTC')
        doc.add_paragraph(f'Rotation: {rotation_name}')
        
        # Add horizontal line
        doc.add_paragraph('_' * 80)
        
        # Get sections in order
        sections = sorted(template.get("sections", []), key=lambda s: s.get("order", 999))
        
        for section in sections:
            section_name = section.get("name")
            
            # Skip if not in include list
            if include_sections and section_name not in include_sections:
                continue
            
            section_title = section.get("title", section_name.replace("_", " ").title())
            section_data = case_data.get("sections", {}).get(section_name, {})
            
            # Section heading
            doc.add_heading(section_title, level=1)
            
            # Add section content
            for field in section.get("fields", []):
                field_name = field.get("name")
                field_label = field.get("label", field_name.replace("_", " ").title())
                field_value = section_data.get("data", {}).get(field_name, "")
                
                # Add field as paragraph
                p = doc.add_paragraph()
                run_label = p.add_run(f'{field_label}: ')
                run_label.bold = True
                
                if field_value:
                    p.add_run(str(field_value))
                else:
                    run_not_doc = p.add_run('Not documented')
                    run_not_doc.italic = True
                    run_not_doc.font.color.rgb = RGBColor(128, 128, 128)
            
            # Add spacing between sections
            doc.add_paragraph()
        
        # Footer
        doc.add_paragraph('_' * 80)
        footer = doc.add_paragraph('Document generated by ClerKase - AI-Powered Clinical Clerking Assistant')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.runs[0].font.size = Pt(8)
        footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        
        return doc
    
    def compile_case_summary(
        self,
        case_id: str,
        case_data: Dict[str, Any]
    ) -> DocumentResult:
        """
        Compile a brief case summary
        
        Args:
            case_id: The case ID
            case_data: Complete case data
            
        Returns:
            DocumentResult with markdown content
        """
        try:
            demographics = case_data.get("sections", {}).get("demographics", {}).get("data", {})
            presenting = case_data.get("sections", {}).get("presenting_complaint", {}).get("data", {})
            
            summary = f"""# Case Summary

**Patient:** {demographics.get('name', 'Unknown')}, {demographics.get('age', '?')} years old, {demographics.get('gender', '?')}  
**Presenting Complaint:** {presenting.get('complaint', 'Not documented')}  
**Duration:** {presenting.get('duration', 'Not documented')}  

**Case ID:** {case_id}  
**Generated:** {datetime.utcnow().strftime('%Y-%m-%d %H:%M')} UTC  

---

This is a brief summary. For the full clerking document, please export the complete case.
"""
            
            return DocumentResult(
                success=True,
                file_path=None,
                content=summary,
                error=None
            )
            
        except Exception as e:
            return DocumentResult(
                success=False,
                file_path=None,
                content=None,
                error=str(e)
            )


# Singleton instance
_compiler = None


def get_document_compiler(output_dir: str = "exports") -> DocumentCompiler:
    """Get or create DocumentCompiler singleton"""
    global _compiler
    if _compiler is None:
        _compiler = DocumentCompiler(output_dir)
    return _compiler
